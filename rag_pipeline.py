import os
import glob

# ── Set HF token before any HuggingFace imports ───────────────────────
try:
    import streamlit as st
    _hf_token = st.secrets.get("HF_TOKEN", "")
    if _hf_token:
        os.environ["HF_TOKEN"] = _hf_token
        os.environ["HUGGINGFACEHUB_API_TOKEN"] = _hf_token
except Exception:
    pass

from langchain_community.document_loaders import TextLoader, PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document

# HuggingFaceEmbeddings moved to langchain-huggingface in langchain v0.3
try:
    from langchain_huggingface import HuggingFaceEmbeddings
except ImportError:
    from langchain_community.embeddings import HuggingFaceEmbeddings  # type: ignore

# ---------- OPTIONAL IMPORTS (graceful fallback) ----------
# OCR is disabled on Streamlit Cloud — too slow for free tier CPU.
OCR_AVAILABLE = False

try:
    import fitz
    import pytesseract
    from PIL import Image
    import io
except ImportError:
    pass

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx unavailable. Install: pip install python-docx")


# ---------- TEXT CLEANING ----------

def clean_text(text: str) -> str:
    """Normalise whitespace while preserving paragraph breaks."""
    lines = [line.strip() for line in str(text).splitlines()]
    cleaned, prev_blank = [], False
    for line in lines:
        if line == "":
            if not prev_blank:
                cleaned.append("")
            prev_blank = True
        else:
            cleaned.append(line)
            prev_blank = False
    return " ".join(cleaned).strip()


# ---------- OCR FALLBACK FOR IMAGE-BASED PDFs ----------

def ocr_pdf(file_path: str) -> list[Document]:
    """
    Rasterise every page of a PDF and run Tesseract OCR on it.
    Used when PyPDFLoader extracts fewer than 50 characters per page.
    """
    if not OCR_AVAILABLE:
        return []

    docs = []
    pdf  = fitz.open(file_path)
    for page_num, page in enumerate(pdf):
        pix  = page.get_pixmap(dpi=300)
        img  = Image.open(io.BytesIO(pix.tobytes("png")))
        text = pytesseract.image_to_string(img)
        if text.strip():
            docs.append(Document(
                page_content=clean_text(text),
                metadata={
                    "source":   file_path,
                    "filename": os.path.basename(file_path),
                    "page":     page_num + 1,
                    "loader":   "ocr",
                }
            ))
    pdf.close()
    return docs


# ---------- LOAD DOCUMENTS ----------

def load_documents(data_path: str = "data") -> list[Document]:
    """Load PDF, TXT, MD and DOCX files from data_path. Skips Excel/CSV."""
    docs = []

    for file_path in glob.glob(f"{data_path}/**/*", recursive=True):
        if not os.path.isfile(file_path):
            continue

        file = os.path.basename(file_path)

        # Skip Excel and CSV files — not processed
        if file.endswith((".xlsx", ".xls", ".csv")):
            print(f"  ⏭  Skipping tabular file: {file}")
            continue

        try:
            # ── TXT / MARKDOWN ──────────────────────────────────────
            if file.endswith((".txt", ".md")):
                loaded = TextLoader(file_path, encoding="utf-8").load()
                for d in loaded:
                    docs.append(Document(
                        page_content=clean_text(d.page_content),
                        metadata={"source": file_path, "filename": file, "loader": "text"}
                    ))

            # ── PDF ─────────────────────────────────────────────────
            elif file.endswith(".pdf"):
                loaded     = PyPDFLoader(file_path).load()
                total_chars = sum(len(d.page_content) for d in loaded)
                avg_chars   = total_chars / max(len(loaded), 1)

                if avg_chars >= 50:
                    for d in loaded:
                        docs.append(Document(
                            page_content=clean_text(d.page_content),
                            metadata={
                                "source":   file_path,
                                "filename": file,
                                "page":     d.metadata.get("page", "?"),
                                "loader":   "pypdf",
                            }
                        ))
                else:
                    print(f"  ↳ Low text yield in {file} — switching to OCR")
                    ocr_docs = ocr_pdf(file_path)
                    if ocr_docs:
                        docs.extend(ocr_docs)
                    else:
                        for d in loaded:
                            docs.append(Document(
                                page_content=clean_text(d.page_content),
                                metadata={"source": file_path, "filename": file, "loader": "pypdf-fallback"}
                            ))

            # ── WORD DOCUMENTS (.docx) ───────────────────────────────
            elif file.endswith(".docx") and DOCX_AVAILABLE:
                word_doc  = DocxDocument(file_path)
                full_text = "\n".join(
                    para.text for para in word_doc.paragraphs if para.text.strip()
                )
                for table in word_doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(
                            cell.text.strip() for cell in row.cells if cell.text.strip()
                        )
                        if row_text:
                            full_text += "\n" + row_text

                if full_text.strip():
                    docs.append(Document(
                        page_content=clean_text(full_text),
                        metadata={"source": file_path, "filename": file, "loader": "docx"}
                    ))

        except Exception as e:
            print(f"⚠️  Error loading {file_path}: {e}")

    print(f"✅  Loaded {len(docs)} document sections from {data_path}/")
    return docs


# ---------- SPLIT DOCUMENTS ----------

def split_documents(docs: list[Document]) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=150,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = splitter.split_documents(docs)

    # Deduplicate by content hash
    seen, unique = set(), []
    for chunk in chunks:
        h = hash(chunk.page_content.strip())
        if h not in seen:
            seen.add(h)
            unique.append(chunk)

    print(f"✅  {len(unique)} unique chunks after splitting & deduplication")
    return unique


# ---------- EMBEDDINGS ----------

def get_embeddings() -> HuggingFaceEmbeddings:
    # Model is committed directly into the repo under models/all-MiniLM-L6-v2/
    # This avoids any HuggingFace network download on Streamlit Cloud.
    base_dir   = os.path.dirname(os.path.abspath(__file__))
    local_path = os.path.join(base_dir, "models", "all-MiniLM-L6-v2")

    if os.path.isdir(local_path):
        model_name = local_path
        print(f"✅  Loading embedding model from local path: {local_path}")
    else:
        model_name = "sentence-transformers/all-MiniLM-L6-v2"
        print(f"⚠️  Local model not found — downloading from HuggingFace")

    embeddings = HuggingFaceEmbeddings(
        model_name=model_name,
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )

    # Validate embeddings loaded correctly
    try:
        test_vec = embeddings.embed_query("test")
        if len(test_vec) < 100:
            raise ValueError(f"Embedding dimension too small: {len(test_vec)}")
        print(f"✅  Embeddings validated — dimension: {len(test_vec)}")
    except Exception as e:
        raise RuntimeError(f"Embedding model failed to load properly: {e}")

    return embeddings


# ---------- VECTOR STORE ----------

def create_vectorstore(chunks: list[Document]) -> FAISS:
    embeddings = get_embeddings()
    return FAISS.from_documents(chunks, embeddings)


# ---------- BUILD INDEX ----------

def build_index(data_path: str = "data") -> FAISS:
    docs   = load_documents(data_path)
    chunks = split_documents(docs)
    print(f"🔨  Building FAISS index over {len(chunks)} chunks …")
    vectorstore = create_vectorstore(chunks)
    vectorstore.save_local("faiss_index")
    print("✅  Index saved to faiss_index/")
    return vectorstore


# ---------- LOAD INDEX ----------

def load_index() -> FAISS:
    if not os.path.exists("faiss_index"):
        print("Index not found. Building new index …")
        return build_index()
    embeddings = get_embeddings()
    return FAISS.load_local(
        "faiss_index",
        embeddings,
        allow_dangerous_deserialization=True,
    )


# ---------- FILTERED RETRIEVAL BY FILE ----------

def get_chunks_from_file(
    vectorstore: FAISS,
    query: str,
    filename: str,
    k: int = 5,
) -> list[Document]:
    """Return the top-k chunks from a specific file matching the query."""
    all_docs = vectorstore.as_retriever(search_kwargs={"k": 20}).invoke(query)
    filtered = [
        d for d in all_docs
        if filename.lower() in d.metadata.get("source", "").lower()
        or filename.lower() in d.metadata.get("filename", "").lower()
    ]
    return filtered[:k]
