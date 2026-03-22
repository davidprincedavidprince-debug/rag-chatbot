"""
index_manager.py

Incremental FAISS index updates.

On Streamlit Cloud the faiss_index/ directory is NOT in Git — it lives
only on the ephemeral container filesystem. This means:

  - First deploy  → full build (index doesn't exist yet)
  - After git push → Streamlit redeploys → full build again
    (ephemeral FS is wiped on each deploy)
  - Within a running session → "Smart sync" uses incremental logic
    (catches any data/ changes pushed mid-session without a redeploy)

Locally the index persists between runs, so Smart sync is genuinely
incremental and saves time.

Manifest format  (faiss_index/manifest.json):
{
  "data/docs/report.pdf":  "sha256hex...",
  "data/mapping.xlsx":     "sha256hex...",
  ...
}
"""

import os
import json
import hashlib
import glob

from langchain_community.vectorstores import FAISS
from rag_pipeline import (
    load_documents,
    split_documents,
    get_embeddings,
    clean_text,
    _ingest_dataframe,
    OCR_AVAILABLE,
    DOCX_AVAILABLE,
)

INDEX_PATH    = "faiss_index"
MANIFEST_PATH = os.path.join(INDEX_PATH, "manifest.json")


# ── Hashing ───────────────────────────────────────────────────────────

def file_hash(path: str) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def scan_data_files(data_path: str = "data") -> dict:
    result = {}
    for path in glob.glob(f"{data_path}/**/*", recursive=True):
        if os.path.isfile(path):
            result[path] = file_hash(path)
    return result


# ── Manifest ──────────────────────────────────────────────────────────

def load_manifest() -> dict:
    if os.path.exists(MANIFEST_PATH):
        with open(MANIFEST_PATH) as f:
            return json.load(f)
    return {}


def save_manifest(manifest: dict) -> None:
    os.makedirs(INDEX_PATH, exist_ok=True)
    with open(MANIFEST_PATH, "w") as f:
        json.dump(manifest, f, indent=2)


# ── Diff ──────────────────────────────────────────────────────────────

def compute_diff(old: dict, new: dict) -> tuple:
    added    = [p for p in new if p not in old]
    modified = [p for p in new if p in old and new[p] != old[p]]
    deleted  = [p for p in old if p not in new]
    return added, modified, deleted


# ── Single-file loader ────────────────────────────────────────────────

def _load_single(file_path: str) -> list:
    """Parse one file using the same loaders as load_documents()."""
    import pandas as pd
    from langchain_community.document_loaders import TextLoader, PyPDFLoader
    from langchain_core.documents import Document

    docs = []
    file = os.path.basename(file_path)

    try:
        if file.endswith((".txt", ".md")):
            for d in TextLoader(file_path, encoding="utf-8").load():
                docs.append(Document(
                    page_content=clean_text(d.page_content),
                    metadata={"source": file_path, "filename": file, "loader": "text"},
                ))

        elif file.endswith(".pdf"):
            from rag_pipeline import ocr_pdf
            loaded    = PyPDFLoader(file_path).load()
            avg_chars = sum(len(d.page_content) for d in loaded) / max(len(loaded), 1)
            if avg_chars >= 50:
                for d in loaded:
                    docs.append(Document(
                        page_content=clean_text(d.page_content),
                        metadata={"source": file_path, "filename": file,
                                  "page": d.metadata.get("page", "?"), "loader": "pypdf"},
                    ))
            else:
                ocr_docs = ocr_pdf(file_path) if OCR_AVAILABLE else []
                docs.extend(ocr_docs or [
                    Document(
                        page_content=clean_text(d.page_content),
                        metadata={"source": file_path, "filename": file, "loader": "pypdf-fallback"},
                    ) for d in loaded
                ])

        elif file.endswith(".docx") and DOCX_AVAILABLE:
            from docx import Document as DocxDoc
            wd        = DocxDoc(file_path)
            full_text = "\n".join(p.text for p in wd.paragraphs if p.text.strip())
            for table in wd.tables:
                for row in table.rows:
                    row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text:
                        full_text += "\n" + row_text
            if full_text.strip():
                docs.append(Document(
                    page_content=clean_text(full_text),
                    metadata={"source": file_path, "filename": file, "loader": "docx"},
                ))

        elif file.endswith(".csv"):
            _ingest_dataframe(pd.read_csv(file_path), file_path, file, docs)

        elif file.endswith((".xlsx", ".xls")):
            xl = pd.ExcelFile(file_path)
            for sheet in xl.sheet_names:
                _ingest_dataframe(xl.parse(sheet), file_path, file, docs, sheet=sheet)

    except Exception as e:
        print(f"  ⚠️  Could not load {file_path}: {e}")

    return docs


# ── Core update function ──────────────────────────────────────────────

def incremental_update(
    data_path: str = "data",
    force_rebuild: bool = False,
) -> tuple:
    """
    Build or incrementally update the FAISS index.

    Returns (vectorstore, stats) where stats contains:
        added, modified, deleted, unchanged, total_files, rebuild (bool)
    """
    embeddings   = get_embeddings()
    old_manifest = load_manifest()
    new_manifest = scan_data_files(data_path)

    added, modified, deleted = compute_diff(old_manifest, new_manifest)

    stats = {
        "added":       len(added),
        "modified":    len(modified),
        "deleted":     len(deleted),
        "unchanged":   len(new_manifest) - len(added) - len(modified),
        "total_files": len(new_manifest),
        "rebuild":     False,
    }

    index_exists = os.path.exists(os.path.join(INDEX_PATH, "index.faiss"))

    # ── Try Drive restore before full rebuild ─────────────────
    if not index_exists and not force_rebuild:
        print("No local index — checking Google Drive …")
        try:
            from hf_store import download_index, hf_index_exists as drive_index_exists
            if drive_index_exists():
                print("Found index on Drive — downloading …")
                if download_index():
                    index_exists = os.path.exists(
                        os.path.join(INDEX_PATH, "index.faiss")
                    )
        except Exception as e:
            print(f"Drive restore skipped: {e}")

    # ── Full rebuild ──────────────────────────────────────────────────
    if force_rebuild or not index_exists:
        reason = "forced" if force_rebuild else "no existing index"
        print(f"Full rebuild ({reason}) — indexing {len(new_manifest)} files …")
        docs   = load_documents(data_path)
        chunks = split_documents(docs)
        vs     = FAISS.from_documents(chunks, embeddings)
        vs.save_local(INDEX_PATH)
        save_manifest(new_manifest)
        try:
            from hf_store import upload_index
            upload_index()
        except Exception as e:
            print(f"Drive upload skipped: {e}")
        stats["rebuild"] = True
        return vs, stats

    # ── Load existing index ───────────────────────────────────────────
    vs = FAISS.load_local(
        INDEX_PATH,
        embeddings,
        allow_dangerous_deserialization=True,
    )

    # ── Nothing changed ───────────────────────────────────────────────
    if not added and not modified and not deleted:
        print("Index is up to date — no changes detected.")
        return vs, stats

    print(f"Incremental update — "
          f"+{len(added)} added  ~{len(modified)} modified  -{len(deleted)} deleted")

    # ── Remove stale chunks ───────────────────────────────────────────
    stale = set(modified + deleted)
    if stale:
        all_ids  = list(vs.docstore._dict.keys())
        keep_ids = [
            i for i in all_ids
            if vs.docstore._dict[i].metadata.get("source", "") not in stale
        ]
        if keep_ids:
            kept_docs = [vs.docstore._dict[i] for i in keep_ids]
            vs = FAISS.from_documents(kept_docs, embeddings)
        else:
            vs = None

    # ── Embed new / changed files ─────────────────────────────────────
    files_to_reindex = added + modified
    if files_to_reindex:
        new_docs = []
        for path in files_to_reindex:
            new_docs.extend(_load_single(path))

        if new_docs:
            new_chunks = split_documents(new_docs)
            print(f"  Embedding {len(new_chunks)} new chunks …")
            if vs is None:
                vs = FAISS.from_documents(new_chunks, embeddings)
            else:
                vs.add_documents(new_chunks)

    # ── Edge case: all files deleted ──────────────────────────────────
    if vs is None:
        from langchain_core.documents import Document
        vs = FAISS.from_documents(
            [Document(page_content="empty index", metadata={"source": "__placeholder__"})],
            embeddings,
        )

    vs.save_local(INDEX_PATH)
    save_manifest(new_manifest)
    return vs, stats
